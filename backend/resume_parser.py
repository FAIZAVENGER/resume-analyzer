import PyPDF2
from docx import Document
import re
import os
from typing import Dict, List, Optional, Tuple
import json
from datetime import datetime
from .nlp_processor import NLPProcessor

class ResumeParser:
    """Advanced resume parser with multiple file format support"""
    
    def __init__(self):
        self.nlp = NLPProcessor()
        
    def parse_file(self, file_path: str) -> Dict:
        """Parse resume file based on extension"""
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == '.pdf':
            text = self.extract_from_pdf(file_path)
        elif ext in ['.docx', '.doc']:
            text = self.extract_from_docx(file_path)
        elif ext == '.txt':
            text = self.extract_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")
        
        return self.analyze_resume(text)
    
    def extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                    except Exception as e:
                        print(f"⚠️ PDF page {page_num + 1} extraction error: {e}")
                        
                # If no text extracted, try alternative method
                if not text.strip():
                    try:
                        # Try reading raw bytes
                        with open(file_path, 'rb') as f:
                            content = f.read()
                            text = content.decode('utf-8', errors='ignore')
                    except:
                        text = "Could not extract text from PDF"
                        
        except Exception as e:
            print(f"❌ PDF parsing error: {e}")
            text = f"Error reading PDF: {str(e)}"
            
        return text
    
    def extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\n"
                        
        except Exception as e:
            print(f"❌ DOCX parsing error: {e}")
            text = f"Error reading DOCX: {str(e)}"
            
        return text
    
    def extract_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            encodings = ['utf-8', 'latin-1', 'windows-1252', 'cp1252', 'utf-16']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                        return text
                except UnicodeDecodeError:
                    continue
                    
            return "Error: Could not decode text file with common encodings"
            
        except Exception as e:
            print(f"❌ TXT parsing error: {e}")
            return f"Error reading TXT: {str(e)}"
    
    def analyze_resume(self, text: str) -> Dict:
        """Comprehensive resume analysis"""
        # Basic cleaning
        text = self.clean_text(text)
        
        # Extract sections
        sections = self.nlp.extract_sections(text)
        
        # Extract entities
        entities = self.nlp.extract_entities(text)
        
        # Extract skills
        skills = self.nlp.extract_skills(text)
        
        # Calculate quality metrics
        quality_metrics = self.nlp.calculate_text_quality(text)
        
        # Extract personal information
        personal_info = self.extract_personal_info(text, entities)
        
        # Extract education details
        education_details = self.extract_education_details(sections.get('education', ''))
        
        # Extract work experience
        work_experience = self.extract_work_experience(sections.get('experience', ''))
        
        # Extract certifications
        certifications = self.extract_certifications(sections.get('certifications', text))
        
        # Extract projects
        projects = self.extract_projects(sections.get('projects', text))
        
        return {
            'raw_text': text[:5000],  # Limit raw text in response
            'sections': sections,
            'entities': entities,
            'skills': skills,
            'quality_metrics': quality_metrics,
            'personal_info': personal_info,
            'education': education_details,
            'experience': work_experience,
            'certifications': certifications,
            'projects': projects,
            'analysis_timestamp': datetime.now().isoformat()
        }
    
    def clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep basic punctuation
        text = re.sub(r'[^\w\s.,;:!?()\-&@+#%/*]', ' ', text)
        
        # Normalize line endings
        text = re.sub(r'\r\n', '\n', text)
        
        return text.strip()
    
    def extract_personal_info(self, text: str, entities: Dict) -> Dict:
        """Extract personal information"""
        info = {
            'name': '',
            'email': '',
            'phone': '',
            'location': '',
            'linkedin': '',
            'portfolio': ''
        }
        
        # Extract name (first PERSON entity)
        if entities['persons']:
            info['name'] = entities['persons'][0]
        
        # Extract email
        if entities['emails']:
            info['email'] = entities['emails'][0]
        
        # Extract phone
        if entities['phones']:
            info['phone'] = entities['phones'][0]
        
        # Extract location
        if entities['locations']:
            info['location'] = entities['locations'][0]
        
        # Extract LinkedIn profile
        linkedin_patterns = [
            r'linkedin\.com/in/([\w\-]+)',
            r'linkedin\.com/pub/([\w\-]+)',
            r'linkedin:?\s*(?:profile:?\s*)?([\w\s\.\-]+)'
        ]
        
        for pattern in linkedin_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                info['linkedin'] = match.group(1)
                break
        
        # Extract portfolio/github
        github_pattern = r'github\.com/([\w\-]+)'
        match = re.search(github_pattern, text, re.IGNORECASE)
        if match:
            info['portfolio'] = f"https://github.com/{match.group(1)}"
        
        return info
    
    def extract_education_details(self, education_text: str) -> List[Dict]:
        """Extract structured education information"""
        education_details = []
        
        # Split into entries
        entries = re.split(r'\n(?=\s*(?:[A-Z][^a-z]*:|[\d]{4}\s*[-–]|(?i)(university|college|institute|school)))', education_text)
        
        for entry in entries:
            if not entry.strip():
                continue
            
            # Extract degree
            degree_pattern = r'(?:Bachelor|Master|PhD|Doctorate|MBA|B\.S\.|B\.A\.|M\.S\.|M\.A\.|Associate|Diploma|Certificate)\s*(?:of\s*(?:Science|Arts|Engineering|Business|Technology))?\s*(?:in\s*[\w\s]+)?'
            degree_match = re.search(degree_pattern, entry, re.IGNORECASE)
            
            # Extract institution
            institution_pattern = r'(?:University|College|Institute|School|Academy)\s+of?\s*[\w\s]+'
            institution_match = re.search(institution_pattern, entry, re.IGNORECASE)
            
            # Extract dates
            date_pattern = r'(\d{4})\s*[-–]\s*(present|\d{4})'
            date_matches = re.findall(date_pattern, entry)
            
            # Extract GPA
            gpa_pattern = r'GPA[:]?\s*([0-4]\.\d{1,2})'
            gpa_match = re.search(gpa_pattern, entry, re.IGNORECASE)
            
            education_details.append({
                'degree': degree_match.group() if degree_match else '',
                'institution': institution_match.group() if institution_match else '',
                'dates': list(date_matches[0]) if date_matches else ['', ''],
                'gpa': gpa_match.group(1) if gpa_match else '',
                'description': entry.strip()
            })
        
        return education_details
    
    def extract_work_experience(self, experience_text: str) -> List[Dict]:
        """Extract structured work experience"""
        experience_details = []
        
        # Split into job entries
        entries = re.split(r'\n(?=\s*(?:[A-Z][^a-z]*:|[\d]{4}\s*[-–]|(?i)(worked|employed|intern|consultant)))', experience_text)
        
        for entry in entries:
            if not entry.strip():
                continue
            
            # Extract job title
            title_pattern = r'^([A-Z][\w\s&]+(?:Manager|Engineer|Developer|Analyst|Specialist|Director|Lead|Head))'
            title_match = re.search(title_pattern, entry)
            
            # Extract company
            company_pattern = r'at\s+([A-Z][\w\s&]+)|,\s*([A-Z][\w\s&]+)'
            company_match = re.search(company_pattern, entry)
            
            # Extract dates
            date_pattern = r'(\d{4})\s*[-–]\s*(present|\d{4})'
            date_matches = re.findall(date_pattern, entry)
            
            # Extract location
            location_pattern = r'[\w\s]+,\s*[A-Z]{2}|Remote|Hybrid|On-site'
            location_match = re.search(location_pattern, entry)
            
            # Extract bullet points
            bullets = re.findall(r'[•\-\*]\s*(.+?)(?=\n[•\-\*]|\n\n|$)', entry, re.DOTALL)
            
            experience_details.append({
                'title': title_match.group(1) if title_match else '',
                'company': company_match.group(1) if company_match else '',
                'dates': list(date_matches[0]) if date_matches else ['', ''],
                'location': location_match.group() if location_match else '',
                'description': entry.strip(),
                'bullet_points': bullets,
                'duration_months': self.calculate_duration(date_matches[0] if date_matches else None)
            })
        
        return experience_details
    
    def extract_certifications(self, cert_text: str) -> List[Dict]:
        """Extract certifications"""
        certifications = []
        
        # Split by common separators
        entries = re.split(r'\n(?=[•\-\*]|\d+\.|\w+\.)', cert_text)
        
        for entry in entries:
            entry = entry.strip()
            if not entry or len(entry) < 5:
                continue
            
            # Look for certification patterns
            cert_patterns = [
                r'([A-Z]+)\s+(?:Certified|Certification)',
                r'(?:AWS|Azure|Google Cloud)\s+([A-Za-z\s]+)',
                r'([A-Za-z\s]+)\s+(?:Certification|Certificate)'
            ]
            
            for pattern in cert_patterns:
                match = re.search(pattern, entry, re.IGNORECASE)
                if match:
                    # Extract date if present
                    date_match = re.search(r'(\d{4})', entry)
                    
                    certifications.append({
                        'name': match.group(1).strip(),
                        'issuer': self.extract_issuer(entry),
                        'date': date_match.group(1) if date_match else '',
                        'description': entry
                    })
                    break
        
        return certifications
    
    def extract_projects(self, projects_text: str) -> List[Dict]:
        """Extract project information"""
        projects = []
        
        # Split by project indicators
        entries = re.split(r'\n(?=(?:Project|App|System|Website|Platform):|[\d]+\.\s*[A-Z])', projects_text)
        
        for entry in entries:
            entry = entry.strip()
            if not entry or len(entry) < 10:
                continue
            
            # Extract project name
            name_match = re.search(r'(?:Project|App|System|Website|Platform):?\s*([^\n]+)', entry)
            
            # Extract technologies
            tech_keywords = ['built with', 'using', 'technologies:', 'stack:', 'tools:']
            technologies = []
            
            for keyword in tech_keywords:
                tech_match = re.search(rf'{keyword}\s*([^\n]+)', entry, re.IGNORECASE)
                if tech_match:
                    tech_list = re.split(r'[,\|]', tech_match.group(1))
                    technologies.extend([t.strip() for t in tech_list])
            
            # Extract URL if present
            url_match = re.search(r'https?://[^\s]+', entry)
            
            projects.append({
                'name': name_match.group(1).strip() if name_match else 'Unnamed Project',
                'description': entry,
                'technologies': list(set(technologies)),  # Remove duplicates
                'url': url_match.group() if url_match else '',
                'features': self.extract_project_features(entry)
            })
        
        return projects
    
    def extract_issuer(self, text: str) -> str:
        """Extract certification issuer"""
        issuers = ['AWS', 'Microsoft', 'Google', 'Cisco', 'Oracle', 'IBM', 'PMI', 'Scrum.org']
        
        for issuer in issuers:
            if issuer.lower() in text.lower():
                return issuer
        
        return 'Unknown'
    
    def extract_project_features(self, text: str) -> List[str]:
        """Extract project features/achievements"""
        features = []
        
        # Look for bullet points or achievements
        bullet_pattern = r'[•\-\*]\s*([^\n]+)'
        bullets = re.findall(bullet_pattern, text)
        
        for bullet in bullets:
            if len(bullet) > 10:  # Minimum length for meaningful feature
                features.append(bullet.strip())
        
        return features
    
    def calculate_duration(self, dates: Optional[Tuple[str, str]]) -> int:
        """Calculate duration in months"""
        if not dates:
            return 0
        
        try:
            start_year = int(dates[0])
            end_year = dates[1]
            
            if end_year.lower() == 'present':
                end_year = datetime.now().year
            
            end_year = int(end_year)
            
            return (end_year - start_year) * 12
            
        except:
            return 0
